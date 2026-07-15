"""
报告生成引擎 — AML 反洗钱风险评估专用
"""
import json
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.database import get_connection


class ReportEngine:

    def _get_aml_data(self, assessment_id: int = None) -> dict:
        """获取AML评估数据"""
        conn = get_connection()
        cursor = conn.cursor()

        if assessment_id:
            cursor.execute("SELECT * FROM assessment_history WHERE id = ?", (assessment_id,))
        else:
            cursor.execute("SELECT * FROM assessment_history ORDER BY assess_date DESC LIMIT 1")
        row = cursor.fetchone()

        if not row:
            conn.close()
            return {}

        r = dict(row)
        dims = json.loads(r["dimensions_json"] or "{}")
        recs = json.loads(r["recommendations_json"] or "[]")

        # Company name from DB or default
        company = "XX基金管理公司"

        # Build details per dimension
        def dim_detail(dk, label):
            dim = dims.get(dk, {})
            items = dim.get("items", [])
            lines = [f"维度：{label} · 得分：{dim.get('score', 0)} 分 · {dim.get('summary', '')}"]
            for item in items:
                risk_tag = {"高": "【高风险】", "中": "【中风险】", "低": "【低风险】"}.get(item.get("risk", ""), "")
                lines.append(f"  · {item.get('name', '')}：{risk_tag} {item.get('detail', '')}")
            return "\n".join(lines)

        # High risk items
        all_items = []
        for dk in ["customer", "product", "channel", "geography"]:
            for item in dims.get(dk, {}).get("items", []):
                if item.get("risk") == "高":
                    all_items.append(f"· [{dk}] {item.get('name', '')} — {item.get('detail', '')}")

        # Build placeholders
        risk_level_names = {"低": "低风险", "中": "中风险", "高": "高风险", "最高": "最高风险"}
        risk_level_name = risk_level_names.get(r["risk_level"], r["risk_level"])

        four_scores = (
            f"- 客户风险维度：{r['customer_score']} 分\n"
            f"- 产品/业务风险维度：{r['product_score']} 分\n"
            f"- 渠道风险维度：{r['channel_score']} 分\n"
            f"- 地域风险维度：{r['geography_score']} 分"
        )

        all_dim_detail = (
            dim_detail("customer", "客户风险") + "\n\n" +
            dim_detail("product", "产品/业务风险") + "\n\n" +
            dim_detail("channel", "渠道风险") + "\n\n" +
            dim_detail("geography", "地域风险")
        )

        conn.close()

        return {
            "公司名称": company,
            "评估日期": r.get("assess_date", ""),
            "报告编号": f"AML-{r.get('assess_date', '')[0:7] if r.get('assess_date') else ''}",
            "编制部门": "风险管理部",
            "评估概要": f"本次评估基于{r.get('preset_id', '基金模板')}，覆盖4个维度24项指标。",
            "综合评分": str(r["overall_score"]),
            "风险等级": risk_level_name,
            "客户数": str(r.get("customer_count", 0)),
            "账户数": str(r.get("account_count", 0)),
            "交易数": str(r.get("trans_count", 0)),
            "产品数": str(r.get("product_count", 0)),
            "四维评分": four_scores,
            "客户评分": str(r.get("customer_score", 0)),
            "产品评分": str(r.get("product_score", 0)),
            "渠道评分": str(r.get("channel_score", 0)),
            "地域评分": str(r.get("geography_score", 0)),
            "客户风险详情": dim_detail("customer", "客户风险"),
            "产品风险详情": dim_detail("product", "产品/业务风险"),
            "渠道风险详情": dim_detail("channel", "渠道风险"),
            "地域风险详情": dim_detail("geography", "地域风险"),
            "全维度详情": all_dim_detail,
            "高风险汇总": "\n".join(all_items) if all_items else "本次评估未发现高风险指标",
            "整改建议": "\n".join(f"{i+1}. {rec}" for i, rec in enumerate(recs)) if recs else "暂无整改建议",
            "结论": f"综合评估{company}洗钱风险等级为{risk_level_name}，"
                     f"综合评分{r['overall_score']}分。建议持续关注高风险指标，并定期复评估。",
        }

    def generate(self, template: dict, info: dict) -> dict:
        """根据模板和参数生成报告"""
        assessment_id = info.get("assessment_id")
        placeholders = self._get_aml_data(assessment_id)

        # Defaults
        placeholders.setdefault("公司名称", info.get("company", "XX基金管理公司"))
        placeholders.setdefault("评估日期", info.get("date", datetime.now().strftime("%Y-%m-%d")))
        placeholders.setdefault("报告编号", info.get("report_no", ""))
        placeholders.setdefault("编制部门", info.get("author", "风险管理部"))

        sections = []
        for sec in template.get("sections", []):
            content = sec.get("content", "")
            for key, val in placeholders.items():
                content = content.replace("{{" + key + "}}", val)
            sections.append({"title": sec["title"], "content": content, "source": "aml_assessment"})

        return {
            "title": info.get("title", ""),
            "report_no": info.get("report_no", ""),
            "author": info.get("author", ""),
            "date": info.get("date", datetime.now().strftime("%Y-%m-%d")),
            "sections": sections,
        }

    def export_word(self, report: dict) -> BytesIO:
        """导出 Word 文档"""
        doc = Document()

        title = doc.add_heading(report.get("title", "报告"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_text = f"报告编号：{report.get('report_no', '')}　　编制人：{report.get('author', '')}　　日期：{report.get('date', '')}"
        info_p = doc.add_paragraph(info_text)
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in info_p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        for sec in report.get("sections", []):
            doc.add_heading(sec["title"], level=1)
            content = sec.get("content", "")
            if content:
                for line in content.split("\n"):
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(11)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
